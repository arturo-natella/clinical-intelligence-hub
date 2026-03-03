"""
Clinical Intelligence Hub — Doctor Visit Prep Generator

Generates a structured visit preparation sheet with:
  1. Active Conditions Summary
  2. Recent Symptoms (with counter-evidence)
  3. Flagged Labs (abnormal with trends)
  4. Questions to Ask
  5. Medication Review (with flagged interactions)
  6. Counter-Evidence Summary (the killer section)
  7. Symptom Patterns (frequency/severity trends)

Can produce on-screen JSON or a formatted Word doc for printing.
"""

import logging
from datetime import datetime, date, timedelta
from pathlib import Path
from typing import Optional

logger = logging.getLogger("CIH-VisitPrep")


class VisitPrepGenerator:
    """Builds a structured visit prep from profile data."""

    def __init__(self, api_key: str = None):
        self._api_key = api_key

    def generate(self, profile_data: dict) -> dict:
        """Return structured visit prep with all sections as JSON."""
        timeline = profile_data.get("clinical_timeline", {})
        analysis = profile_data.get("analysis", {})
        demographics = profile_data.get("demographics", {})

        sections = {
            "generated_at": datetime.now().isoformat(),
            "conditions": self._active_conditions(timeline),
            "recent_symptoms": self._recent_symptoms(timeline),
            "flagged_labs": self._flagged_labs(timeline),
            "questions": self._questions_to_ask(analysis, timeline, profile_data),
            "medications": self._medication_review(timeline, analysis),
            "counter_evidence": self._counter_evidence_summary(timeline),
            "patterns": self._symptom_patterns(timeline),
        }

        # Gemini narrative summary (optional)
        if self._api_key:
            sections["narrative"] = self._generate_narrative(sections)
        else:
            sections["narrative"] = None

        return sections

    # ── Section 1: Active Conditions ─────────────────────

    def _active_conditions(self, timeline: dict) -> list:
        """Current diagnoses with status."""
        diagnoses = timeline.get("diagnoses", [])
        active = []
        for dx in diagnoses:
            status = dx.get("status", "").lower()
            if status in ("active", "current", "ongoing", "chronic", ""):
                active.append({
                    "name": dx.get("name", "Unknown"),
                    "status": dx.get("status", "active"),
                    "date_diagnosed": dx.get("date_diagnosed"),
                    "icd10": dx.get("icd10_code"),
                })
        return active

    # ── Section 2: Recent Symptoms ───────────────────────

    def _recent_symptoms(self, timeline: dict, days: int = 30) -> list:
        """Symptoms from the last N days, sorted by severity."""
        symptoms = timeline.get("symptoms", [])
        cutoff = (date.today() - timedelta(days=days)).isoformat()
        result = []

        severity_order = {"high": 0, "mid": 1, "low": 2}

        for sym in symptoms:
            recent_episodes = []
            for ep in sym.get("episodes", []):
                ep_date = ep.get("episode_date") or ep.get("date_logged", "")
                if isinstance(ep_date, str) and ep_date >= cutoff:
                    recent_episodes.append(ep)

            if recent_episodes:
                # Compute dominant severity
                sev_counts = {"high": 0, "mid": 0, "low": 0}
                for ep in recent_episodes:
                    sev = ep.get("severity", "mid")
                    sev_counts[sev] = sev_counts.get(sev, 0) + 1
                dominant = max(sev_counts, key=sev_counts.get)

                # Collect counter stats
                counter_stats = self._compute_counter_stats_for_symptom(sym)

                result.append({
                    "name": sym.get("symptom_name", "Unknown"),
                    "episode_count": len(recent_episodes),
                    "total_episodes": len(sym.get("episodes", [])),
                    "dominant_severity": dominant,
                    "counter_stats": counter_stats,
                    "latest_episode": recent_episodes[-1] if recent_episodes else None,
                })

        # Sort by severity (high first)
        result.sort(key=lambda x: severity_order.get(x["dominant_severity"], 9))
        return result

    # ── Section 3: Flagged Labs ──────────────────────────

    def _flagged_labs(self, timeline: dict) -> list:
        """Abnormal lab results with trend indicators."""
        labs = timeline.get("labs", [])
        flagged = []

        # Group by lab name
        lab_groups = {}
        for lab in labs:
            name = lab.get("name", "Unknown")
            if name not in lab_groups:
                lab_groups[name] = []
            lab_groups[name].append(lab)

        for name, entries in lab_groups.items():
            # Sort by date (newest first)
            entries.sort(
                key=lambda x: x.get("test_date") or "",
                reverse=True,
            )

            latest = entries[0]
            flag = latest.get("flag", "").lower()

            if flag and flag != "normal":
                trend = "→"  # stable
                if len(entries) >= 2:
                    try:
                        v1 = float(latest.get("value", 0))
                        v2 = float(entries[1].get("value", 0))
                        if v1 > v2 * 1.05:
                            trend = "↑"
                        elif v1 < v2 * 0.95:
                            trend = "↓"
                    except (ValueError, TypeError):
                        pass

                flagged.append({
                    "name": name,
                    "value": latest.get("value"),
                    "unit": latest.get("unit", ""),
                    "flag": flag,
                    "reference_range": latest.get("reference_range", ""),
                    "test_date": latest.get("test_date"),
                    "trend": trend,
                    "history_count": len(entries),
                })

        return flagged

    # ── Section 4: Questions to Ask ──────────────────────

    def _questions_to_ask(
        self, analysis: dict, timeline: dict, profile_data: dict = None,
    ) -> list:
        """Auto-generated questions from flags, interactions, symptoms,
        and missing negative monitoring gaps."""
        questions = []

        # From existing questions_for_doctor
        for q in analysis.get("questions_for_doctor", []):
            questions.append({"source": "analysis", "question": q})

        # From flagged interactions
        for ix in analysis.get("drug_interactions", []):
            sev = ix.get("severity", "").lower()
            if sev in ("high", "critical", "serious"):
                drug_a = ix.get("drug_a", "?")
                drug_b = ix.get("drug_b") or ix.get("gene", "?")
                questions.append({
                    "source": "interaction",
                    "question": (
                        f"I'm taking {drug_a} and {drug_b} together — "
                        f"are there concerns about this combination?"
                    ),
                })

        # From clinical flags
        for flag in analysis.get("flags", []):
            sev = flag.get("severity", "").lower()
            if sev in ("high", "critical"):
                questions.append({
                    "source": "flag",
                    "question": (
                        f"Regarding the flagged finding: {flag.get('title', '')}"
                        f" — what should I know?"
                    ),
                })

        # From counter-evidence
        symptoms = timeline.get("symptoms", [])
        for sym in symptoms:
            for counter in sym.get("counter_definitions", []):
                if counter.get("archived"):
                    continue
                stats = self._single_counter_stats(sym, counter)
                if stats and stats.get("verdict") == "Strongly contradicts":
                    questions.append({
                        "source": "counter_evidence",
                        "question": (
                            f"You've said my {sym.get('symptom_name', '')} "
                            f"are caused by {counter.get('doctor_claim', '')}. "
                            f"My tracked data shows {stats.get('summary', '')}. "
                            f"Could there be another explanation?"
                        ),
                    })

        # From missing negative monitoring gaps
        if profile_data:
            try:
                from src.analysis.missing_negatives import MissingNegativeDetector

                detector = MissingNegativeDetector()
                gaps = detector.analyze(profile_data)

                for gap in gaps:
                    if gap["status"] == "never_tested":
                        questions.append({
                            "source": "monitoring_gap",
                            "question": (
                                f"I have {gap['condition'].lower()} but I don't "
                                f"see a {gap['missing_test']} in my records. "
                                f"Should we schedule one? It's typically done "
                                f"{gap['expected_frequency']}."
                            ),
                        })
                    else:
                        questions.append({
                            "source": "monitoring_gap",
                            "question": (
                                f"My last {gap['missing_test']} for "
                                f"{gap['condition'].lower()} monitoring was "
                                f"~{gap.get('months_overdue', '?')} months "
                                f"overdue. Should we recheck?"
                            ),
                        })
            except Exception as e:
                logger.debug("Missing negatives in visit prep: %s", e)

        return questions

    # ── Section 5: Medication Review ─────────────────────

    def _medication_review(self, timeline: dict, analysis: dict) -> list:
        """Current meds with any flagged interactions."""
        meds = timeline.get("medications", [])
        interactions = analysis.get("drug_interactions", [])

        # Build interaction lookup
        ix_map = {}
        for ix in interactions:
            for drug in [ix.get("drug_a", ""), ix.get("drug_b", "")]:
                drug_lower = drug.lower()
                if drug_lower not in ix_map:
                    ix_map[drug_lower] = []
                ix_map[drug_lower].append({
                    "severity": ix.get("severity"),
                    "description": ix.get("description"),
                })

        result = []
        for med in meds:
            status = med.get("status", "").lower()
            if status not in ("discontinued", "stopped", "completed"):
                name = med.get("name", "Unknown")
                result.append({
                    "name": name,
                    "dosage": med.get("dosage", ""),
                    "frequency": med.get("frequency", ""),
                    "reason": med.get("reason", ""),
                    "interactions": ix_map.get(name.lower(), []),
                })

        return result

    # ── Section 6: Counter-Evidence Summary ──────────────

    def _counter_evidence_summary(self, timeline: dict) -> list:
        """Dedicated section: all counter-evidence with verdicts."""
        symptoms = timeline.get("symptoms", [])
        evidence = []

        for sym in symptoms:
            for counter in sym.get("counter_definitions", []):
                stats = self._single_counter_stats(sym, counter)
                if stats:
                    evidence.append({
                        "symptom": sym.get("symptom_name", ""),
                        "doctor_claim": counter.get("doctor_claim", ""),
                        "measure_type": counter.get("measure_type", ""),
                        "archived": counter.get("archived", False),
                        **stats,
                    })

        return evidence

    # ── Section 7: Symptom Patterns ──────────────────────

    def _symptom_patterns(self, timeline: dict) -> list:
        """Basic frequency and severity trends per symptom."""
        symptoms = timeline.get("symptoms", [])
        patterns = []

        for sym in symptoms:
            episodes = sym.get("episodes", [])
            if not episodes:
                continue

            # Frequency: episodes per week (last 4 weeks)
            four_weeks_ago = (date.today() - timedelta(days=28)).isoformat()
            recent = [
                e for e in episodes
                if (e.get("episode_date") or "") >= four_weeks_ago
            ]
            freq_per_week = round(len(recent) / 4, 1) if recent else 0

            # Severity trend: compare first half vs second half
            sorted_eps = sorted(
                episodes,
                key=lambda e: e.get("episode_date") or "",
            )
            sev_map = {"high": 3, "mid": 2, "low": 1}

            if len(sorted_eps) >= 4:
                mid = len(sorted_eps) // 2
                first_avg = sum(
                    sev_map.get(e.get("severity", "mid"), 2)
                    for e in sorted_eps[:mid]
                ) / mid
                second_avg = sum(
                    sev_map.get(e.get("severity", "mid"), 2)
                    for e in sorted_eps[mid:]
                ) / (len(sorted_eps) - mid)

                if second_avg > first_avg + 0.3:
                    trend = "worsening"
                elif second_avg < first_avg - 0.3:
                    trend = "improving"
                else:
                    trend = "stable"
            else:
                trend = "insufficient_data"

            # Time-of-day clustering
            tod_counts = {}
            for ep in episodes:
                tod = ep.get("time_of_day")
                if tod:
                    tod_counts[tod] = tod_counts.get(tod, 0) + 1
            peak_time = max(tod_counts, key=tod_counts.get) if tod_counts else None

            patterns.append({
                "name": sym.get("symptom_name", ""),
                "total_episodes": len(episodes),
                "freq_per_week": freq_per_week,
                "severity_trend": trend,
                "peak_time_of_day": peak_time,
            })

        return patterns

    # ── Counter-Evidence Helpers ─────────────────────────

    def _compute_counter_stats_for_symptom(self, symptom: dict) -> list:
        """Compute stats for all active counters on a symptom."""
        stats = []
        for counter in symptom.get("counter_definitions", []):
            if counter.get("archived"):
                continue
            s = self._single_counter_stats(symptom, counter)
            if s:
                stats.append(s)
        return stats

    def _single_counter_stats(self, symptom: dict, counter: dict) -> Optional[dict]:
        """Compute stats + verdict for one counter definition."""
        cid = counter.get("counter_id")
        mtype = counter.get("measure_type", "")
        claim = counter.get("doctor_claim", "")
        episodes = symptom.get("episodes", [])

        # Collect values for this counter across all episodes
        values = []
        for ep in episodes:
            cv = ep.get("counter_values", {})
            if cid in cv:
                values.append(cv[cid])

        if not values:
            return None

        result = {
            "doctor_claim": claim,
            "measure_type": mtype,
            "episode_count": len(values),
            "total_episodes": len(episodes),
        }

        if mtype == "scale":
            nums = [v for v in values if isinstance(v, (int, float))]
            if nums:
                avg = round(sum(nums) / len(nums), 1)
                result["average"] = avg
                result["summary"] = (
                    f"Average {claim} level {avg}/5 "
                    f"across {len(nums)} episodes"
                )
                # Verdict
                if avg < 2.0:
                    result["verdict"] = "Strongly contradicts"
                elif avg < 3.0:
                    result["verdict"] = "Weakly contradicts"
                elif avg < 3.5:
                    result["verdict"] = "Inconclusive"
                else:
                    result["verdict"] = "Supports claim"
            else:
                return None

        elif mtype == "yes_no":
            yes_count = sum(1 for v in values if v is True)
            no_count = sum(1 for v in values if v is False)
            total = yes_count + no_count
            if total > 0:
                no_pct = round(no_count / total * 100)
                result["yes_count"] = yes_count
                result["no_count"] = no_count
                result["no_percentage"] = no_pct
                result["summary"] = (
                    f"{claim}: Yes {yes_count}, No {no_count} "
                    f"({no_pct}% No across {total} episodes)"
                )
                if no_pct > 70:
                    result["verdict"] = "Strongly contradicts"
                elif no_pct > 50:
                    result["verdict"] = "Weakly contradicts"
                elif no_pct > 30:
                    result["verdict"] = "Inconclusive"
                else:
                    result["verdict"] = "Supports claim"
            else:
                return None

        elif mtype == "free_text":
            result["entries"] = [str(v) for v in values[:10]]
            result["summary"] = (
                f"{len(values)} free-text entries for '{claim}'"
            )
            result["verdict"] = "Review needed"

        return result

    # ── Gemini Narrative ─────────────────────────────────

    def _generate_narrative(self, sections: dict) -> Optional[str]:
        """Use Gemini to write a plain-English visit prep summary."""
        try:
            import google.generativeai as genai
            import json

            genai.configure(api_key=self._api_key)
            model = genai.GenerativeModel("gemini-2.0-flash")

            prompt = (
                "You are helping a patient prepare for a doctor visit. "
                "Write a concise, empathetic 3-4 paragraph summary they can "
                "read before their appointment. Use plain language (6th grade "
                "reading level). Focus on:\n"
                "1. Key concerns to bring up\n"
                "2. Counter-evidence that challenges doctor assumptions\n"
                "3. Questions to ask\n\n"
                f"Data:\n{json.dumps(sections, default=str)[:6000]}"
            )

            response = model.generate_content(prompt)
            return response.text

        except Exception as e:
            logger.warning("Gemini narrative failed: %s", e)
            return None

    # ── Word Document Export ─────────────────────────────

    def generate_docx(self, sections: dict, output_path: Path) -> Path:
        """Generate a printable Word doc from visit prep sections."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.error("python-docx not installed")
            raise

        doc = Document()

        # Style setup
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # Title
        title = doc.add_heading("Doctor Visit Preparation", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        generated = sections.get("generated_at", "")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Generated: {generated[:10]}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)

        # Narrative (if available)
        narrative = sections.get("narrative")
        if narrative:
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(narrative)

        # Section 1: Active Conditions
        conditions = sections.get("conditions", [])
        if conditions:
            doc.add_heading("Active Conditions", level=1)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Condition"
            hdr[1].text = "Status"
            hdr[2].text = "Since"
            for c in conditions:
                row = table.add_row().cells
                row[0].text = c.get("name", "")
                row[1].text = c.get("status", "")
                row[2].text = c.get("date_diagnosed", "") or "—"

        # Section 2: Recent Symptoms
        symptoms = sections.get("recent_symptoms", [])
        if symptoms:
            doc.add_heading("Recent Symptoms (Last 30 Days)", level=1)
            for sym in symptoms:
                p = doc.add_paragraph()
                run = p.add_run(
                    f"{sym['name']} — {sym['episode_count']} episodes, "
                    f"severity: {sym['dominant_severity'].upper()}"
                )
                run.bold = True

                # Counter stats
                for cs in sym.get("counter_stats", []):
                    summary = cs.get("summary", "")
                    verdict = cs.get("verdict", "")
                    p2 = doc.add_paragraph(style="List Bullet")
                    p2.add_run(f"Doctor says: {cs.get('doctor_claim', '')} → ")
                    run2 = p2.add_run(summary)
                    run2.font.color.rgb = RGBColor(200, 160, 50)
                    if verdict:
                        p2.add_run(f"  [{verdict}]")

        # Section 3: Flagged Labs
        flagged = sections.get("flagged_labs", [])
        if flagged:
            doc.add_heading("Flagged Lab Results", level=1)
            table = doc.add_table(rows=1, cols=5)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Test"
            hdr[1].text = "Value"
            hdr[2].text = "Flag"
            hdr[3].text = "Trend"
            hdr[4].text = "Date"
            for lab in flagged:
                row = table.add_row().cells
                row[0].text = lab.get("name", "")
                val = lab.get("value", "")
                unit = lab.get("unit", "")
                row[1].text = f"{val} {unit}".strip()
                row[2].text = lab.get("flag", "")
                row[3].text = lab.get("trend", "→")
                row[4].text = lab.get("test_date", "") or "—"

        # Section 4: Medication Review
        meds = sections.get("medications", [])
        if meds:
            doc.add_heading("Current Medications", level=1)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Medication"
            hdr[1].text = "Dosage"
            hdr[2].text = "Frequency"
            hdr[3].text = "For"
            for med in meds:
                row = table.add_row().cells
                row[0].text = med.get("name", "")
                row[1].text = med.get("dosage", "")
                row[2].text = med.get("frequency", "")
                row[3].text = med.get("reason", "")

                # Flag interactions
                for ix in med.get("interactions", []):
                    p = doc.add_paragraph(style="List Bullet")
                    run = p.add_run(
                        f"⚠ Interaction: {ix.get('description', '')}"
                    )
                    run.font.color.rgb = RGBColor(200, 80, 60)

        # Section 5: Counter-Evidence Summary (THE KILLER SECTION)
        evidence = sections.get("counter_evidence", [])
        if evidence:
            doc.add_heading("Counter-Evidence Summary", level=1)
            p = doc.add_paragraph(
                "The following data tracks claims your doctor has made "
                "about symptom causes, measured against your actual experience."
            )
            p.runs[0].font.italic = True
            p.runs[0].font.size = Pt(10)

            for ev in evidence:
                status = " (archived)" if ev.get("archived") else ""
                p = doc.add_paragraph()
                run = p.add_run(
                    f"{ev.get('symptom', '')} — Doctor says: "
                    f"{ev.get('doctor_claim', '')}{status}"
                )
                run.bold = True

                summary = ev.get("summary", "")
                verdict = ev.get("verdict", "")
                p2 = doc.add_paragraph(style="List Bullet")
                run2 = p2.add_run(f"Your data: {summary}")
                run2.font.color.rgb = RGBColor(200, 160, 50)

                if verdict:
                    p3 = doc.add_paragraph(style="List Bullet")
                    color = (
                        RGBColor(60, 160, 60)
                        if "contradicts" in verdict.lower()
                        else RGBColor(200, 160, 50)
                        if verdict == "Inconclusive"
                        else RGBColor(200, 80, 60)
                    )
                    run3 = p3.add_run(f"Verdict: {verdict}")
                    run3.bold = True
                    run3.font.color.rgb = color

        # Section 6: Questions to Ask
        questions = sections.get("questions", [])
        if questions:
            doc.add_heading("Questions to Ask Your Doctor", level=1)
            for i, q in enumerate(questions, 1):
                source_label = {
                    "counter_evidence": "Based on your tracked data",
                    "interaction": "Drug interaction concern",
                    "flag": "Flagged finding",
                    "analysis": "From your analysis",
                }.get(q.get("source", ""), "")

                p = doc.add_paragraph(style="List Number")
                p.add_run(q.get("question", ""))
                if source_label:
                    run = p.add_run(f"\n  ({source_label})")
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(128, 128, 128)

        # Section 7: Symptom Patterns
        patterns = sections.get("patterns", [])
        if patterns:
            doc.add_heading("Symptom Patterns", level=1)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Symptom"
            hdr[1].text = "Frequency"
            hdr[2].text = "Trend"
            hdr[3].text = "Peak Time"
            for pat in patterns:
                row = table.add_row().cells
                row[0].text = pat.get("name", "")
                row[1].text = f"{pat.get('freq_per_week', 0)}/week"
                trend_label = {
                    "worsening": "↑ Worsening",
                    "improving": "↓ Improving",
                    "stable": "→ Stable",
                    "insufficient_data": "— Too few episodes",
                }.get(pat.get("severity_trend", ""), "—")
                row[2].text = trend_label
                row[3].text = (pat.get("peak_time_of_day") or "—").capitalize()

        # Disclaimer
        doc.add_paragraph("")
        p = doc.add_paragraph(
            "This document was generated by Clinical Intelligence Hub "
            "for personal use. It is not a medical diagnosis. Always discuss "
            "findings with your healthcare provider."
        )
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        doc.save(str(output_path))
        logger.info("Visit prep saved to %s", output_path)
        return output_path
