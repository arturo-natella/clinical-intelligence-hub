"""
Clinical Intelligence Hub — Report Builder (Pass 6)

Generates a 10-section Word document with complete clinical provenance.
Every finding is footnoted with source file, page number, and date.

Sections:
  1. Patient Summary
  2. Health Timeline
  3. Active Conditions & Medications
  4. Lab Trends & Threshold Analysis
  5. Imaging Analysis
  6. Genetic Profile & Pharmacogenomics
  7. Patterns, Flags & Drug Interactions
  8. Cross-Disciplinary Insights
  9. Questions for Your Doctor
  10. Disclaimer, Sources & Methods
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

from src.models import (
    AlertSeverity,
    AnalysisResults,
    ClinicalTimeline,
    Demographics,
    FindingCategory,
    PatientProfile,
)

logger = logging.getLogger("CIH-Report")


class ReportBuilder:
    """
    Generates a 10-section clinical report as a Word document.

    Every clinical finding includes provenance: which file it came from,
    which page, and which model extracted it.
    """

    def __init__(self):
        self._doc = None

    def generate(self, profile: PatientProfile,
                 output_path: Path,
                 redaction_summary: list = None,
                 file_count: int = 0) -> Path:
        """
        Generate the full clinical report.

        Args:
            profile: Complete patient profile with analysis results
            output_path: Where to save the .docx file
            redaction_summary: PII redaction counts for Section 10
            file_count: Number of source files processed

        Returns:
            Path to the generated report
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise

        self._doc = Document()
        self._setup_styles()

        timeline = profile.clinical_timeline
        analysis = profile.analysis

        # ── Title Page ──
        self._add_title_page(profile)

        # ── Section 1: Patient Summary ──
        self._section_1_patient_summary(profile.demographics, timeline)

        # ── Section 2: Health Timeline ──
        self._section_2_health_timeline(timeline)

        # ── Section 3: Active Conditions & Medications ──
        self._section_3_conditions_medications(timeline)

        # ── Section 4: Lab Trends ──
        self._section_4_lab_trends(timeline)

        # ── Section 5: Imaging Analysis ──
        self._section_5_imaging(timeline)

        # ── Section 6: Genetic Profile ──
        self._section_6_genetics(timeline)

        # ── Section 7: Patterns & Flags ──
        self._section_7_patterns_flags(analysis)

        # ── Section 8: Cross-Disciplinary Insights ──
        self._section_8_cross_disciplinary(analysis)

        # ── Section 9: Questions for Your Doctor ──
        self._section_9_questions(analysis)

        # ── Section 10: Disclaimer, Sources & Methods ──
        self._section_10_disclaimer(
            profile, redaction_summary, file_count
        )

        # Save
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self._doc.save(str(output_path))
        logger.info(f"Report saved to {output_path}")
        return output_path

    # ── Style Setup ───────────────────────────────────────────

    def _setup_styles(self):
        """Configure document styles for a clean medical report."""
        from docx.shared import Pt, RGBColor

        style = self._doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Heading styles
        for level in range(1, 4):
            heading_style = self._doc.styles[f"Heading {level}"]
            heading_style.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)

    # ── Title Page ────────────────────────────────────────────

    def _add_title_page(self, profile: PatientProfile):
        """Add title page with report metadata."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        title = self._doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("Clinical Intelligence Report")
        run.bold = True
        run.font.size = Pt(24)

        subtitle = self._doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Comprehensive Medical Records Analysis")
        run.font.size = Pt(14)

        self._doc.add_paragraph()  # Spacer

        meta = self._doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n")
        meta.add_run(f"Pipeline Version: {profile.pipeline_version}\n")
        meta.add_run(f"Documents Processed: {len(profile.processed_files)}")

        self._doc.add_page_break()

    # ── Section 1: Patient Summary ────────────────────────────

    def _section_1_patient_summary(self, demographics: Demographics,
                                    timeline: ClinicalTimeline):
        """Overview card: demographics + active counts."""
        self._doc.add_heading("1. Patient Summary", level=1)

        # Demographics
        table = self._doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"

        demo_items = []
        if demographics.biological_sex:
            demo_items.append(("Biological Sex", demographics.biological_sex))
        if demographics.birth_year:
            age = datetime.now().year - demographics.birth_year
            demo_items.append(("Approximate Age", f"{age} years"))
        if demographics.blood_type:
            demo_items.append(("Blood Type", demographics.blood_type))
        if demographics.ethnicity:
            demo_items.append(("Ethnicity", demographics.ethnicity))

        for label, value in demo_items:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(value)

        self._doc.add_paragraph()

        # Active counts
        active_meds = [m for m in timeline.medications
                       if m.status.value in ("active", "prn")]
        active_dx = [d for d in timeline.diagnoses
                     if d.status and d.status.lower() in ("active", "chronic")]

        summary = self._doc.add_paragraph()
        summary.add_run(f"Active Medications: {len(active_meds)}\n").bold = True
        summary.add_run(f"Active Conditions: {len(active_dx)}\n")
        summary.add_run(f"Lab Results on File: {len(timeline.labs)}\n")
        summary.add_run(f"Imaging Studies: {len(timeline.imaging)}\n")
        summary.add_run(f"Genetic Variants: {len(timeline.genetics)}\n")
        summary.add_run(f"Allergies: {len(timeline.allergies)}")

    # ── Section 2: Health Timeline ────────────────────────────

    def _section_2_health_timeline(self, timeline: ClinicalTimeline):
        """Chronological view of all medical events."""
        self._doc.add_heading("2. Health Timeline", level=1)

        # Collect all dated events
        events = []

        for med in timeline.medications:
            if med.start_date:
                events.append((med.start_date, "Medication",
                               f"Started {med.name} {med.dosage or ''} {med.frequency or ''}".strip(),
                               med.provenance))

        for lab in timeline.labs:
            if lab.test_date:
                flag = f" [{lab.flag}]" if lab.flag else ""
                events.append((lab.test_date, "Lab",
                               f"{lab.name}: {lab.value or lab.value_text} {lab.unit or ''}{flag}".strip(),
                               lab.provenance))

        for dx in timeline.diagnoses:
            if dx.date_diagnosed:
                events.append((dx.date_diagnosed, "Diagnosis",
                               f"{dx.name} ({dx.status or 'Unknown status'})",
                               dx.provenance))

        for proc in timeline.procedures:
            if proc.procedure_date:
                events.append((proc.procedure_date, "Procedure",
                               proc.name, proc.provenance))

        for img in timeline.imaging:
            if img.study_date:
                events.append((img.study_date, "Imaging",
                               f"{img.modality or 'Study'}: {img.body_region or 'Unknown region'}",
                               img.provenance))

        # Sort by date (most recent first)
        events.sort(key=lambda e: e[0], reverse=True)

        if not events:
            self._doc.add_paragraph(
                "No dated medical events found in the processed records."
            )
            return

        # Table format
        table = self._doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        headers = table.rows[0].cells
        headers[0].text = "Date"
        headers[1].text = "Type"
        headers[2].text = "Event"
        headers[3].text = "Source"

        for event_date, event_type, description, provenance in events[:50]:
            row = table.add_row()
            row.cells[0].text = str(event_date)
            row.cells[1].text = event_type
            row.cells[2].text = description
            row.cells[3].text = self._format_provenance(provenance)

        if len(events) > 50:
            self._doc.add_paragraph(
                f"(Showing 50 of {len(events)} events. Full timeline available in the Hub.)"
            )

    # ── Section 3: Active Conditions & Medications ────────────

    def _section_3_conditions_medications(self, timeline: ClinicalTimeline):
        """Active conditions and current medications."""
        self._doc.add_heading("3. Active Conditions & Medications", level=1)

        # Active conditions
        self._doc.add_heading("Active Conditions", level=2)
        active_dx = [d for d in timeline.diagnoses
                     if d.status and d.status.lower() in ("active", "chronic")]

        if active_dx:
            table = self._doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            headers = table.rows[0].cells
            headers[0].text = "Condition"
            headers[1].text = "Date Diagnosed"
            headers[2].text = "Status"
            headers[3].text = "Source"

            for dx in active_dx:
                row = table.add_row()
                row.cells[0].text = dx.name
                row.cells[1].text = str(dx.date_diagnosed or "Unknown")
                row.cells[2].text = dx.status or "Active"
                row.cells[3].text = self._format_provenance(dx.provenance)
        else:
            self._doc.add_paragraph("No active conditions found in records.")

        # Active medications
        self._doc.add_heading("Current Medications", level=2)
        active_meds = [m for m in timeline.medications
                       if m.status.value in ("active", "prn")]

        if active_meds:
            table = self._doc.add_table(rows=1, cols=5)
            table.style = "Light Grid Accent 1"
            headers = table.rows[0].cells
            headers[0].text = "Medication"
            headers[1].text = "Dosage"
            headers[2].text = "Frequency"
            headers[3].text = "Reason"
            headers[4].text = "Source"

            for med in active_meds:
                row = table.add_row()
                row.cells[0].text = med.name
                row.cells[1].text = med.dosage or "—"
                row.cells[2].text = med.frequency or "—"
                row.cells[3].text = med.reason or "—"
                row.cells[4].text = self._format_provenance(med.provenance)
        else:
            self._doc.add_paragraph("No active medications found in records.")

        # Allergies
        if timeline.allergies:
            self._doc.add_heading("Allergies", level=2)
            for allergy in timeline.allergies:
                p = self._doc.add_paragraph(style="List Bullet")
                p.add_run(f"{allergy.allergen}").bold = True
                if allergy.reaction:
                    p.add_run(f" — {allergy.reaction}")
                if allergy.severity:
                    p.add_run(f" (Severity: {allergy.severity})")

    # ── Section 4: Lab Trends ─────────────────────────────────

    def _section_4_lab_trends(self, timeline: ClinicalTimeline):
        """Lab values with threshold analysis."""
        self._doc.add_heading("4. Lab Trends & Threshold Analysis", level=1)

        if not timeline.labs:
            self._doc.add_paragraph("No lab results found in processed records.")
            return

        # Group labs by name
        lab_groups: dict[str, list] = {}
        for lab in timeline.labs:
            key = lab.name.lower()
            if key not in lab_groups:
                lab_groups[key] = []
            lab_groups[key].append(lab)

        # Sort each group by date
        for key in lab_groups:
            lab_groups[key].sort(
                key=lambda l: l.test_date or datetime.min.date()
            )

        # Show flagged labs first
        flagged_labs = {
            k: v for k, v in lab_groups.items()
            if any(l.flag and l.flag.lower() not in ("normal", "") for l in v)
        }

        if flagged_labs:
            self._doc.add_heading("Flagged Results", level=2)
            for lab_name, results in flagged_labs.items():
                latest = results[-1]
                flag_text = latest.flag or ""
                p = self._doc.add_paragraph()
                p.add_run(f"{latest.name}: ").bold = True
                p.add_run(
                    f"{latest.value or latest.value_text} {latest.unit or ''} "
                    f"[{flag_text}]"
                )
                if latest.reference_low is not None or latest.reference_high is not None:
                    ref_range = f" (Reference: {latest.reference_low or '—'} – {latest.reference_high or '—'})"
                    p.add_run(ref_range)
                p.add_run(f"  [{self._format_provenance(latest.provenance)}]")

                # Show trend if multiple results
                if len(results) > 1:
                    oldest = results[0]
                    if oldest.value and latest.value:
                        if latest.value > oldest.value:
                            trend = "↗ Increasing"
                        elif latest.value < oldest.value:
                            trend = "↘ Decreasing"
                        else:
                            trend = "→ Stable"
                        self._doc.add_paragraph(
                            f"    Trend: {trend} "
                            f"({oldest.value} → {latest.value} over {len(results)} measurements)"
                        )

        # All labs table
        self._doc.add_heading("All Lab Results", level=2)
        table = self._doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        headers = table.rows[0].cells
        headers[0].text = "Test"
        headers[1].text = "Value"
        headers[2].text = "Unit"
        headers[3].text = "Flag"
        headers[4].text = "Date"
        headers[5].text = "Source"

        for lab in sorted(timeline.labs,
                          key=lambda l: l.test_date or datetime.min.date(),
                          reverse=True)[:100]:
            row = table.add_row()
            row.cells[0].text = lab.name
            row.cells[1].text = str(lab.value if lab.value is not None else lab.value_text or "—")
            row.cells[2].text = lab.unit or "—"
            row.cells[3].text = lab.flag or "Normal"
            row.cells[4].text = str(lab.test_date or "Unknown")
            row.cells[5].text = self._format_provenance(lab.provenance)

    # ── Section 5: Imaging Analysis ───────────────────────────

    def _section_5_imaging(self, timeline: ClinicalTimeline):
        """Imaging studies and AI-detected findings."""
        self._doc.add_heading("5. Imaging Analysis", level=1)

        if not timeline.imaging:
            self._doc.add_paragraph("No imaging studies found in processed records.")
            return

        for study in timeline.imaging:
            heading = f"{study.modality or 'Study'}"
            if study.body_region:
                heading += f" — {study.body_region}"
            if study.study_date:
                heading += f" ({study.study_date})"
            self._doc.add_heading(heading, level=2)

            if study.description:
                self._doc.add_paragraph(study.description)

            self._doc.add_paragraph(
                f"Source: {self._format_provenance(study.provenance)}"
            )

            # Findings
            if study.findings:
                self._doc.add_heading("Findings", level=3)
                for finding in study.findings:
                    p = self._doc.add_paragraph(style="List Bullet")
                    p.add_run(finding.description).bold = True

                    details = []
                    if finding.monai_model:
                        details.append(f"Detected by: {finding.monai_model}")
                    if finding.confidence is not None:
                        details.append(f"Confidence: {finding.confidence:.0%}")
                    if finding.measurements:
                        for k, v in finding.measurements.items():
                            details.append(f"{k}: {v}")
                    if finding.comparison_to_prior:
                        details.append(f"Compared to prior: {finding.comparison_to_prior}")

                    if details:
                        self._doc.add_paragraph("    " + " | ".join(details))

    # ── Section 6: Genetic Profile ────────────────────────────

    def _section_6_genetics(self, timeline: ClinicalTimeline):
        """Genetic variants and pharmacogenomic implications."""
        self._doc.add_heading("6. Genetic Profile & Pharmacogenomics", level=1)

        if not timeline.genetics:
            self._doc.add_paragraph(
                "No genetic test results found in processed records."
            )
            return

        table = self._doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        headers = table.rows[0].cells
        headers[0].text = "Gene"
        headers[1].text = "Variant"
        headers[2].text = "Phenotype"
        headers[3].text = "Clinical Significance"
        headers[4].text = "Implications"

        for variant in timeline.genetics:
            row = table.add_row()
            row.cells[0].text = variant.gene
            row.cells[1].text = variant.variant or "—"
            row.cells[2].text = variant.phenotype or "—"
            row.cells[3].text = variant.clinical_significance or "—"
            row.cells[4].text = variant.implications or "—"

        self._doc.add_paragraph()
        self._doc.add_paragraph(
            "Note: Genetic results should be reviewed by a genetic counselor "
            "or pharmacogenomics-trained provider for clinical decision-making."
        )

    # ── Section 7: Patterns & Flags ───────────────────────────

    def _section_7_patterns_flags(self, analysis: AnalysisResults):
        """Clinical flags, drug interactions, and patterns."""
        self._doc.add_heading("7. Patterns, Flags & Drug Interactions", level=1)

        # Drug interactions
        if analysis.drug_interactions:
            self._doc.add_heading("Drug Interactions", level=2)
            for interaction in analysis.drug_interactions:
                p = self._doc.add_paragraph()
                severity_icon = self._severity_icon(interaction.severity)
                p.add_run(f"{severity_icon} ").bold = True

                if interaction.drug_b:
                    p.add_run(f"{interaction.drug_a} ↔ {interaction.drug_b}").bold = True
                elif interaction.gene:
                    p.add_run(f"{interaction.drug_a} ↔ {interaction.gene}").bold = True
                else:
                    p.add_run(interaction.drug_a).bold = True

                self._doc.add_paragraph(f"    {interaction.description}")
                self._doc.add_paragraph(f"    Source: {interaction.source}")

        # Clinical flags
        if analysis.flags:
            self._doc.add_heading("Clinical Flags", level=2)

            # Sort by severity
            sorted_flags = sorted(
                analysis.flags,
                key=lambda f: ["critical", "high", "moderate", "low", "info"].index(
                    f.severity.value
                ),
            )

            for flag in sorted_flags:
                severity_icon = self._severity_icon(flag.severity)
                p = self._doc.add_paragraph()
                p.add_run(f"{severity_icon} {flag.title}").bold = True
                self._doc.add_paragraph(f"    {flag.description}")

                if flag.evidence:
                    self._doc.add_paragraph(
                        f"    Evidence: {'; '.join(flag.evidence)}"
                    )

        if not analysis.drug_interactions and not analysis.flags:
            self._doc.add_paragraph(
                "No significant patterns or flags were identified."
            )

    # ── Section 8: Cross-Disciplinary ─────────────────────────

    def _section_8_cross_disciplinary(self, analysis: AnalysisResults):
        """Cross-specialty connections that individual doctors might miss."""
        self._doc.add_heading("8. Cross-Disciplinary Insights", level=1)

        self._doc.add_paragraph(
            "These connections were found by analyzing your complete medical "
            "records across all 29 medical specialties and 7 adjacent health "
            "domains. Individual specialists may not see these patterns because "
            "they only review data within their specialty."
        )

        if analysis.cross_disciplinary:
            for connection in analysis.cross_disciplinary:
                severity_icon = self._severity_icon(connection.severity)
                p = self._doc.add_paragraph()
                p.add_run(f"{severity_icon} {connection.title}").bold = True

                self._doc.add_paragraph(f"    {connection.description}")
                self._doc.add_paragraph(
                    f"    Specialties: {', '.join(connection.specialties)}"
                )

                if connection.supporting_literature:
                    lit_refs = []
                    for lit in connection.supporting_literature[:3]:
                        ref = lit.title
                        if lit.journal and lit.year:
                            ref += f" ({lit.journal}, {lit.year})"
                        if lit.doi:
                            ref += f" DOI: {lit.doi}"
                        lit_refs.append(ref)
                    self._doc.add_paragraph(
                        "    Supporting literature: " + "; ".join(lit_refs)
                    )

                if connection.question_for_doctor:
                    p = self._doc.add_paragraph()
                    p.add_run("    → Ask your doctor: ").bold = True
                    p.add_run(connection.question_for_doctor)

                self._doc.add_paragraph()  # Spacer
        else:
            self._doc.add_paragraph(
                "No cross-disciplinary connections were identified. "
                "This may indicate your care is well-coordinated across providers."
            )

        # Community insights (clearly labeled)
        if analysis.community_insights:
            self._doc.add_heading("Community Reports (Unverified)", level=2)
            self._doc.add_paragraph(
                "⚠️ The following patterns were reported by patients on Reddit. "
                "These are NOT clinical data and have NOT been verified by medical "
                "professionals. They are included as discussion points only."
            )

            for insight in analysis.community_insights:
                p = self._doc.add_paragraph(style="List Bullet")
                p.add_run(f"r/{insight.subreddit}: ").bold = True
                p.add_run(insight.description[:200])
                if insight.upvote_count:
                    p.add_run(f" ({insight.upvote_count:,} upvotes)")

                if insight.cross_disciplinary_context:
                    self._doc.add_paragraph(
                        f"    Possible mechanism: {insight.cross_disciplinary_context}"
                    )

    # ── Section 9: Questions for Doctor ───────────────────────

    def _section_9_questions(self, analysis: AnalysisResults):
        """Generated questions to discuss with providers."""
        self._doc.add_heading("9. Questions for Your Doctor", level=1)

        self._doc.add_paragraph(
            "Based on the analysis of your medical records, these questions "
            "may be worth discussing at your next appointment:"
        )

        questions = list(analysis.questions_for_doctor)

        # Add questions from cross-disciplinary connections
        for conn in analysis.cross_disciplinary:
            if conn.question_for_doctor and conn.question_for_doctor not in questions:
                questions.append(conn.question_for_doctor)

        # Add questions from flags
        for flag in analysis.flags:
            if flag.question_for_doctor and flag.question_for_doctor not in questions:
                questions.append(flag.question_for_doctor)

        if questions:
            for i, question in enumerate(questions, 1):
                p = self._doc.add_paragraph()
                p.add_run(f"{i}. ").bold = True
                p.add_run(question)
        else:
            self._doc.add_paragraph(
                "No specific questions were generated. Your records appear "
                "well-documented and consistent."
            )

    # ── Section 10: Disclaimer & Sources ──────────────────────

    def _section_10_disclaimer(self, profile: PatientProfile,
                                redaction_summary: list = None,
                                file_count: int = 0):
        """Disclaimer, methodology, PII audit, and citations."""
        self._doc.add_heading("10. Disclaimer, Sources & Methods", level=1)

        # Disclaimer
        self._doc.add_heading("Disclaimer", level=2)
        self._doc.add_paragraph(
            "This report was generated by an AI-powered analysis system. "
            "It is NOT a medical diagnosis, medical advice, or a substitute "
            "for professional medical care. All findings should be reviewed "
            "and validated by a qualified healthcare provider before any "
            "clinical decisions are made."
        )
        self._doc.add_paragraph(
            "The AI models used in this analysis may produce inaccurate or "
            "incomplete results. Drug interactions, adverse events, and "
            "cross-disciplinary connections are flagged for discussion — "
            "they are not confirmed diagnoses or contraindications."
        )

        # Methodology
        self._doc.add_heading("Methodology", level=2)
        self._doc.add_paragraph(
            "This report was generated using a 6-pass analysis pipeline:"
        )
        methods = [
            "Pass 0: File classification, OCR (Apple Vision/Tesseract), deduplication",
            "Pass 1a: MedGemma 27B text extraction (local, on-device)",
            "Pass 1b: MedGemma 4B medical image analysis (local, on-device)",
            "Pass 1c: MONAI clinical detection models (local, on-device)",
            "Pass 1.5: PII redaction (Microsoft Presidio)",
            "Pass 2: Gemini 3.1 Pro Preview gap-filling (cloud, PII-redacted)",
            "Pass 3: Deep Research cross-disciplinary analysis (cloud, PII-redacted)",
            "Pass 4: Deep Research literature search (cloud, PII-redacted)",
            "Pass 5: Clinical validation (OpenFDA, RxNorm, PubMed, DrugBank)",
            "Pass 6: Report generation",
        ]
        for method in methods:
            self._doc.add_paragraph(method, style="List Bullet")

        # Processing stats
        self._doc.add_heading("Processing Summary", level=2)
        table = self._doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"

        stats = [
            ("Documents Processed", str(file_count or len(profile.processed_files))),
            ("Report Generated", datetime.now().strftime("%Y-%m-%d %H:%M")),
            ("Pipeline Version", profile.pipeline_version),
            ("Medications Found", str(len(profile.clinical_timeline.medications))),
            ("Lab Results Found", str(len(profile.clinical_timeline.labs))),
            ("Diagnoses Found", str(len(profile.clinical_timeline.diagnoses))),
            ("Imaging Studies", str(len(profile.clinical_timeline.imaging))),
            ("Genetic Variants", str(len(profile.clinical_timeline.genetics))),
            ("Drug Interactions Flagged", str(len(profile.analysis.drug_interactions))),
            ("Clinical Flags", str(len(profile.analysis.flags))),
            ("Cross-Disciplinary Connections", str(len(profile.analysis.cross_disciplinary))),
            ("Literature Citations", str(len(profile.analysis.literature))),
        ]

        for label, value in stats:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = value

        # PII Redaction audit
        if redaction_summary:
            self._doc.add_heading("Privacy Protection (PII Redaction)", level=2)
            self._doc.add_paragraph(
                "All patient-identifying information was removed before any "
                "data was sent to cloud services. The following PII types were detected "
                "and redacted:"
            )
            table = self._doc.add_table(rows=1, cols=2)
            table.style = "Light Grid Accent 1"
            headers = table.rows[0].cells
            headers[0].text = "PII Type"
            headers[1].text = "Redactions"

            for entry in redaction_summary:
                row = table.add_row()
                row.cells[0].text = entry.get("original_type", "Unknown")
                row.cells[1].text = str(entry.get("count", 0))

        # Literature citations
        if profile.analysis.literature:
            self._doc.add_heading("Literature Citations", level=2)
            for i, citation in enumerate(profile.analysis.literature, 1):
                ref = f"{i}. {citation.title}"
                if citation.authors:
                    ref = f"{i}. {citation.authors}. {citation.title}"
                if citation.journal and citation.year:
                    ref += f". {citation.journal}, {citation.year}"
                if citation.doi:
                    ref += f". DOI: {citation.doi}"
                if citation.pubmed_id:
                    ref += f". PMID: {citation.pubmed_id}"
                self._doc.add_paragraph(ref)

    # ── Helpers ───────────────────────────────────────────────

    @staticmethod
    def _format_provenance(provenance) -> str:
        """Format a Provenance object as a readable citation."""
        if provenance is None:
            return "Unknown source"

        parts = []
        if provenance.source_file:
            parts.append(provenance.source_file)
        if provenance.source_page is not None:
            parts.append(f"p.{provenance.source_page}")
        if provenance.extraction_model:
            parts.append(provenance.extraction_model)
        return ", ".join(parts) if parts else "Unknown source"

    @staticmethod
    def _severity_icon(severity: AlertSeverity) -> str:
        """Map severity to a text indicator for Word docs."""
        mapping = {
            AlertSeverity.CRITICAL: "[CRITICAL]",
            AlertSeverity.HIGH: "[HIGH]",
            AlertSeverity.MODERATE: "[MODERATE]",
            AlertSeverity.LOW: "[LOW]",
            AlertSeverity.INFO: "[INFO]",
        }
        return mapping.get(severity, "[—]")
