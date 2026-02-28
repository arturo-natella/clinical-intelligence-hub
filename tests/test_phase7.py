"""
Phase 7 Tests — Report Generation (Pass 6)

Tests report builder and addendum builder structure and logic.
Actual .docx generation requires python-docx, so these tests
verify both the logic and the actual document generation.
"""

import sys
import tempfile
from datetime import date, datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from src.models import (
    AlertSeverity,
    AnalysisResults,
    ClinicalFlag,
    ClinicalTimeline,
    CrossDisciplinaryConnection,
    Demographics,
    Diagnosis,
    DrugInteraction,
    FindingCategory,
    GeneticVariant,
    ImagingFinding,
    ImagingStudy,
    LabResult,
    LiteratureCitation,
    Medication,
    MedicationStatus,
    MonitoringAlert,
    PatientProfile,
    Procedure,
    Provenance,
)


def _sample_profile() -> PatientProfile:
    """Create a sample patient profile for testing."""
    prov = Provenance(
        source_file="test_records.pdf",
        source_page=1,
        extraction_model="medgemma-27b",
        confidence=0.95,
    )

    return PatientProfile(
        demographics=Demographics(
            biological_sex="Female",
            birth_year=1963,
            blood_type="A+",
        ),
        clinical_timeline=ClinicalTimeline(
            medications=[
                Medication(
                    name="Metformin", dosage="500mg", frequency="twice daily",
                    status=MedicationStatus.ACTIVE, reason="Type 2 Diabetes",
                    start_date=date(2020, 3, 15), provenance=prov,
                ),
                Medication(
                    name="Lisinopril", dosage="10mg", frequency="daily",
                    status=MedicationStatus.ACTIVE, reason="Hypertension",
                    start_date=date(2019, 6, 1), provenance=prov,
                ),
                Medication(
                    name="Atorvastatin", dosage="20mg", frequency="daily",
                    status=MedicationStatus.ACTIVE, reason="High cholesterol",
                    start_date=date(2021, 1, 10), provenance=prov,
                ),
            ],
            labs=[
                LabResult(
                    name="HbA1c", value=7.2, unit="%", flag="High",
                    test_date=date(2024, 6, 15), provenance=prov,
                ),
                LabResult(
                    name="Vitamin D", value=18.0, unit="ng/mL", flag="Low",
                    test_date=date(2024, 6, 15), provenance=prov,
                ),
                LabResult(
                    name="Total Cholesterol", value=195.0, unit="mg/dL",
                    flag="Normal", test_date=date(2024, 6, 15), provenance=prov,
                ),
            ],
            diagnoses=[
                Diagnosis(
                    name="Type 2 Diabetes", status="Active",
                    date_diagnosed=date(2020, 3, 1), provenance=prov,
                ),
                Diagnosis(
                    name="Essential Hypertension", status="Active",
                    date_diagnosed=date(2019, 5, 15), provenance=prov,
                ),
            ],
            imaging=[
                ImagingStudy(
                    modality="CT", body_region="Chest",
                    study_date=date(2024, 1, 20),
                    description="CT Chest without contrast",
                    findings=[
                        ImagingFinding(
                            description="3mm ground glass nodule in RLL",
                            body_region="Right Lower Lobe",
                            measurements={"diameter_mm": 3},
                            confidence=0.87,
                        )
                    ],
                    provenance=prov,
                ),
            ],
            genetics=[
                GeneticVariant(
                    gene="CYP2D6", variant="*4/*4",
                    phenotype="Poor Metabolizer",
                    clinical_significance="Actionable",
                    implications="Reduced metabolism of CYP2D6 substrates",
                    provenance=prov,
                ),
            ],
            procedures=[
                Procedure(
                    name="Colonoscopy", procedure_date=date(2023, 9, 10),
                    provider="Dr. Smith", provenance=prov,
                ),
            ],
        ),
        analysis=AnalysisResults(
            drug_interactions=[
                DrugInteraction(
                    drug_a="Metformin", drug_b="Lisinopril",
                    severity=AlertSeverity.LOW,
                    description="Low risk interaction — monitor renal function",
                    source="RxNorm Interaction API",
                ),
            ],
            flags=[
                ClinicalFlag(
                    category=FindingCategory.LAB_THRESHOLD,
                    severity=AlertSeverity.HIGH,
                    title="Vitamin D deficiency",
                    description="Vitamin D level of 18 ng/mL is deficient (<30).",
                    evidence=["Lab: Vitamin D 18 ng/mL, test_records.pdf p.1"],
                ),
            ],
            cross_disciplinary=[
                CrossDisciplinaryConnection(
                    title="Vitamin D deficiency may contribute to diabetes control",
                    description=(
                        "Multiple studies link vitamin D deficiency to increased "
                        "insulin resistance and poorer glycemic control in T2DM."
                    ),
                    specialties=["Endocrinology", "Nutrition Science"],
                    patient_data_points=["Vitamin D: 18 ng/mL", "HbA1c: 7.2%"],
                    severity=AlertSeverity.MODERATE,
                    question_for_doctor=(
                        "Could vitamin D supplementation help with blood sugar control?"
                    ),
                ),
            ],
            literature=[
                LiteratureCitation(
                    title="Vitamin D and Type 2 Diabetes: A Systematic Review",
                    authors="Smith J et al.",
                    journal="Diabetes Care",
                    year=2023,
                    doi="10.1234/dc.2023.0001",
                    pubmed_id="12345678",
                ),
            ],
            questions_for_doctor=[
                "My HbA1c has been above target — should we adjust Metformin dosage?",
                "Should I be tested for vitamin D deficiency and consider supplementation?",
            ],
        ),
    )


# ── ReportBuilder Tests ──────────────────────────────────────

def test_report_builder_generates_docx():
    """ReportBuilder generates a .docx file."""
    try:
        from docx import Document
    except ImportError:
        print("⊘ Report test skipped — python-docx not installed")
        return

    from src.report.builder import ReportBuilder

    profile = _sample_profile()
    builder = ReportBuilder()

    with tempfile.TemporaryDirectory() as tmpdir:
        output = Path(tmpdir) / "test_report.docx"
        result = builder.generate(
            profile, output,
            redaction_summary=[
                {"original_type": "PERSON", "count": 5},
                {"original_type": "DATE_OF_BIRTH", "count": 2},
            ],
            file_count=3,
        )

        assert result.exists(), "Report file should exist"
        assert result.suffix == ".docx"
        assert result.stat().st_size > 0, "Report should not be empty"

        # Verify it's a valid docx
        doc = Document(str(result))
        full_text = "\n".join([p.text for p in doc.paragraphs])

        assert "Clinical Intelligence Report" in full_text
        assert "Patient Summary" in full_text
        assert "Health Timeline" in full_text
        assert "Active Conditions" in full_text
        assert "Lab Trends" in full_text
        assert "Imaging Analysis" in full_text
        assert "Genetic Profile" in full_text
        assert "Patterns, Flags" in full_text
        assert "Cross-Disciplinary" in full_text
        assert "Questions for Your Doctor" in full_text
        assert "Disclaimer" in full_text

    print("✓ ReportBuilder generates valid .docx with all 10 sections")


def test_report_contains_patient_data():
    """Report includes actual patient data from the profile."""
    try:
        from docx import Document
    except ImportError:
        print("⊘ Report data test skipped — python-docx not installed")
        return

    from src.report.builder import ReportBuilder

    profile = _sample_profile()
    builder = ReportBuilder()

    with tempfile.TemporaryDirectory() as tmpdir:
        output = Path(tmpdir) / "test_report.docx"
        builder.generate(profile, output)

        doc = Document(str(output))
        # Collect text from both paragraphs and table cells
        all_text_parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text_parts.append(cell.text)
        full_text = "\n".join(all_text_parts)

        # Patient data should appear
        assert "Metformin" in full_text, "Should include medication name"
        assert "Type 2 Diabetes" in full_text, "Should include diagnosis"
        assert "Vitamin D" in full_text, "Should include lab name"
        assert "CYP2D6" in full_text, "Should include genetic variant"

    print("✓ Report contains actual patient data")


def test_report_contains_provenance():
    """Report includes source provenance citations."""
    try:
        from docx import Document
    except ImportError:
        print("⊘ Provenance test skipped — python-docx not installed")
        return

    from src.report.builder import ReportBuilder

    profile = _sample_profile()
    builder = ReportBuilder()

    with tempfile.TemporaryDirectory() as tmpdir:
        output = Path(tmpdir) / "test_report.docx"
        builder.generate(profile, output)

        doc = Document(str(output))
        all_text_parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text_parts.append(cell.text)
        full_text = "\n".join(all_text_parts)

        assert "test_records.pdf" in full_text, "Should cite source file"
        assert "medgemma-27b" in full_text, "Should cite extraction model"

    print("✓ Report includes provenance citations")


def test_report_empty_profile():
    """Report handles empty profile gracefully."""
    try:
        from docx import Document
    except ImportError:
        print("⊘ Empty profile test skipped — python-docx not installed")
        return

    from src.report.builder import ReportBuilder

    profile = PatientProfile()
    builder = ReportBuilder()

    with tempfile.TemporaryDirectory() as tmpdir:
        output = Path(tmpdir) / "empty_report.docx"
        result = builder.generate(profile, output)

        assert result.exists()
        doc = Document(str(result))
        assert len(doc.paragraphs) > 10, "Even empty report should have structure"

    print("✓ Report handles empty profile without errors")


def test_report_provenance_formatting():
    """Provenance helper formats correctly."""
    from src.report.builder import ReportBuilder

    prov = Provenance(
        source_file="labs_2024.pdf",
        source_page=3,
        extraction_model="medgemma-27b",
    )

    formatted = ReportBuilder._format_provenance(prov)
    assert "labs_2024.pdf" in formatted
    assert "p.3" in formatted
    assert "medgemma-27b" in formatted

    # None provenance
    formatted = ReportBuilder._format_provenance(None)
    assert "Unknown" in formatted

    print("✓ Provenance formatting works correctly")


def test_report_severity_icons():
    """Severity mapping produces correct labels."""
    from src.report.builder import ReportBuilder

    assert "[CRITICAL]" in ReportBuilder._severity_icon(AlertSeverity.CRITICAL)
    assert "[HIGH]" in ReportBuilder._severity_icon(AlertSeverity.HIGH)
    assert "[MODERATE]" in ReportBuilder._severity_icon(AlertSeverity.MODERATE)
    assert "[LOW]" in ReportBuilder._severity_icon(AlertSeverity.LOW)
    assert "[INFO]" in ReportBuilder._severity_icon(AlertSeverity.INFO)

    print("✓ Severity icons map correctly")


# ── AddendumBuilder Tests ────────────────────────────────────

def test_addendum_generates_docx():
    """AddendumBuilder generates a .docx file."""
    try:
        from docx import Document
    except ImportError:
        print("⊘ Addendum test skipped — python-docx not installed")
        return

    from src.report.addendum import AddendumBuilder

    alert = MonitoringAlert(
        source="OpenFDA",
        title="New safety alert for Metformin",
        description="FDA updated Metformin label with new lactic acidosis guidance.",
        relevance_explanation=(
            "Patient is currently taking Metformin 500mg twice daily. "
            "The new guidance may affect dosing recommendations."
        ),
        severity=AlertSeverity.HIGH,
        url="https://www.fda.gov/example",
    )

    profile = _sample_profile()
    builder = AddendumBuilder()

    with tempfile.TemporaryDirectory() as tmpdir:
        output = Path(tmpdir) / "test_addendum.docx"
        result = builder.generate(alert, profile, output)

        assert result.exists()
        assert result.stat().st_size > 0

        doc = Document(str(result))
        full_text = "\n".join([p.text for p in doc.paragraphs])

        assert "Addendum" in full_text
        assert "Metformin" in full_text
        assert "HIGH" in full_text
        assert "DISCLAIMER" in full_text

    print("✓ AddendumBuilder generates valid addendum .docx")


def test_addendum_batch():
    """AddendumBuilder generates multiple addendums."""
    try:
        from docx import Document
    except ImportError:
        print("⊘ Batch addendum test skipped — python-docx not installed")
        return

    from src.report.addendum import AddendumBuilder

    alerts = [
        MonitoringAlert(
            source="PubMed",
            title="New study on Metformin",
            description="New evidence for Metformin.",
            relevance_explanation="Relevant to patient.",
            severity=AlertSeverity.MODERATE,
        ),
        MonitoringAlert(
            source="OpenFDA",
            title="Drug recall notice",
            description="Recall for Atorvastatin batch.",
            relevance_explanation="Patient takes Atorvastatin.",
            severity=AlertSeverity.HIGH,
        ),
    ]

    profile = _sample_profile()
    builder = AddendumBuilder()

    with tempfile.TemporaryDirectory() as tmpdir:
        output_dir = Path(tmpdir) / "addendums"
        paths = builder.generate_batch(alerts, profile, output_dir)

        assert len(paths) == 2
        for p in paths:
            assert p.exists()
            assert p.suffix == ".docx"

    print("✓ Batch addendum generation works")


# ── Run All Tests ────────────────────────────────────────────

if __name__ == "__main__":
    print("=" * 60)
    print("Phase 7 Tests — Report Generation")
    print("=" * 60)

    # Report builder
    test_report_builder_generates_docx()
    test_report_contains_patient_data()
    test_report_contains_provenance()
    test_report_empty_profile()
    test_report_provenance_formatting()
    test_report_severity_icons()

    # Addendum builder
    test_addendum_generates_docx()
    test_addendum_batch()

    print()
    print("=" * 60)
    print("All Phase 7 tests passed ✓")
    print("=" * 60)
